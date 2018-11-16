#! Python3

import docx

def readDocx():
    doc = docx.Document('demo.docx')
    print(len(doc.paragraphs))
    print(doc.paragraphs[0].text)
    print(doc.paragraphs[1].text)
    print(doc.paragraphs[1].runs)
    print(len(doc.paragraphs[1].runs))
readDocx()